"""
EXTRATOR DE DOCUMENTOS - VERSÃO WEB
Acesso online via navegador
"""

import streamlit as st
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import PyPDF2
from pdf2image import convert_from_bytes
import tempfile

# Configuração da página
st.set_page_config(
    page_title="Extrator de Documentos",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilo CSS customizado
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
        padding: 0.5rem;
        font-size: 16px;
        border-radius: 5px;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-radius: 5px;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        background-color: #d1ecf1;
        border-radius: 5px;
        border-left: 4px solid #17a2b8;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Configuração do Tesseract para ambiente Linux (Streamlit Cloud)
import platform
import os

if platform.system() == "Linux":
    # Streamlit Cloud usa Linux
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"
else:
    # Windows local
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

class ExtratorWeb:
    def __init__(self):
        pass
    
    def preprocessar_imagem(self, img):
        """Pré-processa imagem para OCR"""
        img = img.convert('L')
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2)
        img = img.filter(ImageFilter.SHARPEN)
        
        if img.width < 1000:
            factor = 1000 / img.width
            novo_tamanho = (int(img.width * factor), int(img.height * factor))
            img = img.resize(novo_tamanho, Image.Resampling.LANCZOS)
        
        return img
    
    def extrair_texto_imagem(self, img, preprocessar=True):
        """Extrai texto de imagem"""
        if preprocessar:
            img = self.preprocessar_imagem(img)
        
        config_tesseract = '--psm 6 --oem 3'
        texto = pytesseract.image_to_string(img, lang='por', config=config_tesseract)
        return texto
    
    def extrair_texto_pdf(self, pdf_bytes):
        """Extrai texto de PDF"""
        texto_completo = ""
        
        try:
            pdf_file = io.BytesIO(pdf_bytes)
            leitor = PyPDF2.PdfReader(pdf_file)
            
            for pagina in leitor.pages:
                texto = pagina.extract_text()
                if texto.strip():
                    texto_completo += texto + "\n"
            
            if len(texto_completo.strip()) < 100:
                # PDF escaneado - usa OCR
                imagens = convert_from_bytes(pdf_bytes, dpi=300)
                for imagem in imagens:
                    texto = self.extrair_texto_imagem(imagem)
                    texto_completo += texto + "\n"
                    
        except Exception as e:
            st.error(f"Erro ao processar PDF: {str(e)}")
        
        return texto_completo
    
    def extrair_campo(self, texto, palavras_chave):
        """Extrai campo específico"""
        linhas = texto.split("\n")
        
        for i, linha in enumerate(linhas):
            linha_lower = linha.lower().strip()
            
            for palavra in palavras_chave:
                if palavra.lower() in linha_lower:
                    partes = linha.split(':', 1)
                    if len(partes) > 1:
                        valor = partes[1].strip()
                        if valor:
                            return valor
                    if i + 1 < len(linhas):
                        return linhas[i + 1].strip()
        
        return "NÃO ENCONTRADO"
    
    def extrair_nome(self, texto):
        """Extrai nome"""
        palavras_chave = ['nome:', 'nome completo:', 'empregado:', 'funcionário:']
        nome = self.extrair_campo(texto, palavras_chave)
        
        if nome != "NÃO ENCONTRADO":
            nome = re.sub(r'^[^a-záàâãéèêíïóôõöúçñA-ZÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ]+', '', nome)
            nome = re.split(r'\d{3}\.\d{3}\.\d{3}', nome)[0].strip()
        
        return nome if nome and len(nome) > 3 else "NÃO ENCONTRADO"
    
    def extrair_cpf(self, texto):
        """Extrai CPF"""
        cpf_regex1 = r"\d{3}\.\d{3}\.\d{3}-\d{2}"
        encontrado = re.search(cpf_regex1, texto)
        if encontrado:
            return encontrado.group()
        
        cpf_regex2 = r"\b\d{11}\b"
        encontrado = re.search(cpf_regex2, texto)
        if encontrado:
            cpf = encontrado.group()
            return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
        
        return "NÃO ENCONTRADO"
    
    def extrair_rg(self, texto):
        """Extrai RG com múltiplos padrões"""
        linhas = texto.split("\n")
        
        # Procura por linha com "RG:"
        for i, linha in enumerate(linhas):
            linha_lower = linha.lower().strip()
            
            if 'rg:' in linha_lower or 'rg ' in linha_lower or 'identidade:' in linha_lower:
                # Remove CPF se estiver junto
                linha_limpa = re.sub(r'\d{3}\.\d{3}\.\d{3}-\d{2}', '', linha)
                
                # Tenta extrair da mesma linha
                partes = linha_limpa.split(':', 1)
                if len(partes) > 1:
                    rg_candidato = partes[1].strip()
                    
                    # Procura por padrões de RG
                    # Padrão 1: números com pontos/hífens (ex: 14122330700680943)
                    match = re.search(r'[\d.-]{10,20}', rg_candidato)
                    if match:
                        rg = match.group()
                        apenas_numeros = re.sub(r'[^0-9]', '', rg)
                        if 7 <= len(apenas_numeros) <= 15:
                            return rg
                
                # Tenta próxima linha
                if i + 1 < len(linhas):
                    proxima_linha = linhas[i + 1].strip()
                    match = re.search(r'[\d.-]{10,20}', proxima_linha)
                    if match:
                        rg = match.group()
                        apenas_numeros = re.sub(r'[^0-9]', '', rg)
                        if 7 <= len(apenas_numeros) <= 15:
                            return rg
        
        # Busca alternativa: procura sequências longas de números (que não sejam CPF)
        # RG geralmente tem mais de 11 dígitos quando não formatado
        numeros_longos = re.findall(r'\b\d{12,17}\b', texto)
        if numeros_longos:
            return numeros_longos[0]
        
        return "NÃO ENCONTRADO"
    
    def extrair_endereco(self, texto):
        """Extrai endereço"""
        palavras_chave = ['endereço:', 'residência:', 'endereço residencial:']
        return self.extrair_campo(texto, palavras_chave)
    
    def extrair_funcao(self, texto):
        """Extrai função"""
        palavras_chave = ['função:', 'cargo:', 'ocupação:']
        return self.extrair_campo(texto, palavras_chave)
    
    def extrair_salario(self, texto):
        """Extrai salário"""
        salario_regex = r"R\$\s*\d{1,3}(?:\.\d{3})*(?:,\d{2})?"
        encontrado = re.search(salario_regex, texto)
        if encontrado:
            return encontrado.group()
        
        return "NÃO ENCONTRADO"
    
    def extrair_data_nascimento(self, texto):
        """Extrai data de nascimento com múltiplos padrões"""
        palavras_chave = [
            'data de nascimento:', 'nascimento:', 'data nasc:', 
            'dt. nascimento:', 'dt nascimento:', 'data de nasc:',
            'nascido em:', 'dt. nasc:'
        ]
        
        linhas = texto.split("\n")
        
        # Procura por palavras-chave
        for i, linha in enumerate(linhas):
            linha_lower = linha.lower().strip()
            
            for palavra in palavras_chave:
                if palavra in linha_lower:
                    # Procura data na mesma linha
                    datas_na_linha = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', linha)
                    if datas_na_linha:
                        # Verifica se é uma data de nascimento válida (não muito recente)
                        for data in datas_na_linha:
                            ano = int(data[-4:])
                            if 1920 <= ano <= 2010:
                                return data.replace('-', '/')
                    
                    # Tenta próxima linha
                    if i + 1 < len(linhas):
                        proxima = linhas[i + 1].strip()
                        datas_proxima = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', proxima)
                        if datas_proxima:
                            for data in datas_proxima:
                                ano = int(data[-4:])
                                if 1920 <= ano <= 2010:
                                    return data.replace('-', '/')
        
        # Busca alternativa: procura TODAS as datas e filtra por ano
        todas_datas = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', texto)
        for data in todas_datas:
            try:
                ano = int(data[-4:])
                # Data de nascimento provável: entre 1920 e 2010
                if 1920 <= ano <= 2010:
                    return data.replace('-', '/')
            except:
                continue
        
        return "NÃO ENCONTRADO"
    
    def extrair_data_inicio(self, texto):
        """Extrai data de início/admissão com múltiplos padrões"""
        palavras_chave = [
            'data de início:', 'data de admissão:', 'admissão:', 
            'início:', 'inicio:', 'data de inicio:', 'admitido em:',
            'data admissão:', 'dt. admissão:', 'dt. inicio:',
            'data início:'
        ]
        
        linhas = texto.split("\n")
        
        # Procura por palavras-chave
        for i, linha in enumerate(linhas):
            linha_lower = linha.lower().strip()
            
            for palavra in palavras_chave:
                if palavra in linha_lower:
                    # Procura data na mesma linha
                    datas_na_linha = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', linha)
                    if datas_na_linha:
                        # Data de admissão geralmente é recente (após 1980)
                        for data in datas_na_linha:
                            ano = int(data[-4:])
                            if 1980 <= ano <= 2025:
                                return data.replace('-', '/')
                    
                    # Tenta próxima linha
                    if i + 1 < len(linhas):
                        proxima = linhas[i + 1].strip()
                        datas_proxima = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', proxima)
                        if datas_proxima:
                            for data in datas_proxima:
                                ano = int(data[-4:])
                                if 1980 <= ano <= 2025:
                                    return data.replace('-', '/')
        
        # Busca alternativa: procura datas recentes (provavelmente admissão)
        todas_datas = re.findall(r'\b\d{2}[/-]\d{2}[/-]\d{4}\b', texto)
        for data in todas_datas:
            try:
                ano = int(data[-4:])
                # Data de início mais provável: 1980-2025
                if 2000 <= ano <= 2025:  # Prioriza datas mais recentes
                    return data.replace('-', '/')
            except:
                continue
        
        # Se não achou, aceita datas de 1980+
        for data in todas_datas:
            try:
                ano = int(data[-4:])
                if 1980 <= ano <= 2025:
                    return data.replace('-', '/')
            except:
                continue
        
        return "NÃO ENCONTRADO"
    
    def extrair_dados_completos(self, texto):
        """Extrai todos os dados"""
        return {
            'nome': self.extrair_nome(texto),
            'nacionalidade': 'Brasileiro',
            'data_nascimento': self.extrair_data_nascimento(texto),
            'endereco': self.extrair_endereco(texto),
            'cpf': self.extrair_cpf(texto),
            'rg': self.extrair_rg(texto),
            'funcao': self.extrair_funcao(texto),
            'salario': self.extrair_salario(texto),
            'data_inicio': self.extrair_data_inicio(texto)
        }
    
    def gerar_documento_word(self, dados_list):
        """Gera documento Word"""
        doc = Document()
        
        titulo = doc.add_heading("Relatório de Extração de Dados", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = doc.add_paragraph()
        info.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n").bold = True
        info.add_run(f"Total de registros: {len(dados_list)}\n").bold = True
        
        doc.add_paragraph("_" * 100)
        
        for idx, dados in enumerate(dados_list, 1):
            doc.add_heading(f"Registro {idx}", level=2)
            
            tabela = doc.add_table(rows=9, cols=2)
            tabela.style = 'Light Grid Accent 1'
            
            campos = [
                ("Nome:", dados['nome']),
                ("Nacionalidade:", dados['nacionalidade']),
                ("Data de Nascimento:", dados['data_nascimento']),
                ("Endereço:", dados['endereco']),
                ("CPF:", dados['cpf']),
                ("RG:", dados['rg']),
                ("Função:", dados['funcao']),
                ("Salário:", dados['salario']),
                ("Data de Início:", dados['data_inicio'])
            ]
            
            for i, (campo, valor) in enumerate(campos):
                tabela.rows[i].cells[0].text = campo
                tabela.rows[i].cells[1].text = valor
                tabela.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
        
        # Salva em memória
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io

# Inicializa o extrator
if 'extrator' not in st.session_state:
    st.session_state.extrator = ExtratorWeb()
    st.session_state.dados_extraidos = []

# Interface Principal
st.title("📄 Extrator de Documentos Online")
st.markdown("### Extraia informações de documentos automaticamente")

# Sidebar
with st.sidebar:
    st.header("⚙️ Configurações")
    
    preprocessar = st.checkbox("Pré-processar imagens", value=True, 
                               help="Melhora a qualidade da extração")
    
    st.markdown("---")
    
    st.header("📊 Campos Extraídos")
    st.markdown("""
    - ✅ Nome
    - ✅ Nacionalidade
    - ✅ Data de Nascimento
    - ✅ Endereço
    - ✅ CPF
    - ✅ RG
    - ✅ Função
    - ✅ Salário
    - ✅ Data de Início
    """)
    
    st.markdown("---")
    
    if st.button("🗑️ Limpar Todos os Dados"):
        st.session_state.dados_extraidos = []
        st.success("Dados limpos!")
        st.rerun()

# Tabs principais
tab1, tab2, tab3 = st.tabs(["📤 Upload", "📋 Dados Extraídos", "📝 Gerar Relatório"])

with tab1:
    st.header("📤 Upload de Documentos")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("🖼️ Imagens")
        imagens = st.file_uploader(
            "Selecione imagens",
            type=['jpg', 'jpeg', 'png', 'bmp'],
            accept_multiple_files=True,
            key="imagens"
        )
        
        if imagens:
            if st.button("Processar Imagens", key="btn_img"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, imagem in enumerate(imagens):
                    status_text.text(f"Processando {idx+1}/{len(imagens)}: {imagem.name}")
                    
                    img = Image.open(imagem)
                    texto = st.session_state.extrator.extrair_texto_imagem(img, preprocessar)
                    dados = st.session_state.extrator.extrair_dados_completos(texto)
                    dados['arquivo'] = imagem.name
                    
                    st.session_state.dados_extraidos.append(dados)
                    
                    progress_bar.progress((idx + 1) / len(imagens))
                
                status_text.empty()
                progress_bar.empty()
                st.success(f"✅ {len(imagens)} imagem(ns) processada(s)!")
                st.rerun()
    
    with col2:
        st.subheader("📑 PDFs")
        pdfs = st.file_uploader(
            "Selecione PDFs",
            type=['pdf'],
            accept_multiple_files=True,
            key="pdfs"
        )
        
        if pdfs:
            if st.button("Processar PDFs", key="btn_pdf"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, pdf in enumerate(pdfs):
                    status_text.text(f"Processando {idx+1}/{len(pdfs)}: {pdf.name}")
                    
                    pdf_bytes = pdf.read()
                    texto = st.session_state.extrator.extrair_texto_pdf(pdf_bytes)
                    dados = st.session_state.extrator.extrair_dados_completos(texto)
                    dados['arquivo'] = pdf.name
                    
                    st.session_state.dados_extraidos.append(dados)
                    
                    progress_bar.progress((idx + 1) / len(pdfs))
                
                status_text.empty()
                progress_bar.empty()
                st.success(f"✅ {len(pdfs)} PDF(s) processado(s)!")
                st.rerun()

with tab2:
    st.header("📋 Dados Extraídos")
    
    if not st.session_state.dados_extraidos:
        st.info("👆 Faça upload de documentos na aba 'Upload' para começar")
    else:
        st.success(f"✅ {len(st.session_state.dados_extraidos)} registro(s) extraído(s)")
        
        for idx, dados in enumerate(st.session_state.dados_extraidos, 1):
            with st.expander(f"📄 Registro {idx} - {dados['arquivo']}"):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown(f"**👤 Nome:** {dados['nome']}")
                    st.markdown(f"**🌍 Nacionalidade:** {dados['nacionalidade']}")
                    st.markdown(f"**🎂 Data Nascimento:** {dados['data_nascimento']}")
                    st.markdown(f"**🏠 Endereço:** {dados['endereco']}")
                    st.markdown(f"**🆔 CPF:** {dados['cpf']}")
                
                with col2:
                    st.markdown(f"**📇 RG:** {dados['rg']}")
                    st.markdown(f"**💼 Função:** {dados['funcao']}")
                    st.markdown(f"**💰 Salário:** {dados['salario']}")
                    st.markdown(f"**📅 Data Início:** {dados['data_inicio']}")

with tab3:
    st.header("📝 Gerar Relatório")
    
    if not st.session_state.dados_extraidos:
        st.warning("⚠️ Nenhum dado extraído ainda!")
    else:
        st.info(f"📊 Pronto para gerar relatório com {len(st.session_state.dados_extraidos)} registro(s)")
        
        if st.button("📥 Baixar Relatório Word", type="primary"):
            with st.spinner("Gerando documento..."):
                doc_io = st.session_state.extrator.gerar_documento_word(st.session_state.dados_extraidos)
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nome_arquivo = f"relatorio_extracao_{timestamp}.docx"
                
                st.download_button(
                    label="💾 Download Relatório",
                    data=doc_io,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Relatório gerado com sucesso!")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray;'>
    <p>Desenvolvido por "J@iron Sousa" usando Python + Streamlit</p>
    <p>v1.0 - Dezembro 2024</p>
</div>
""", unsafe_allow_html=True)
